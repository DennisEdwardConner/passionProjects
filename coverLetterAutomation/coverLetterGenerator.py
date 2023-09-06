import re
import os
import lxml.etree
from docx import Document

# Replace these with your actual template, job description, and resume file paths
template_path = os.path.abspath(r'C:\Users\keirn\Desktop\meritamerica\repos\passion projects\coverLetterAutomation\cover_letter_template.docx')
job_description_path = r'C:\Users\keirn\Desktop\meritamerica\repos\passion projects\coverLetterAutomation\job_description.docx'
resume_path = r'C:\Users\keirn\Desktop\meritamerica\repos\passion projects\coverLetterAutomation\resume.docx'

# Read the cover letter template
document = Document(template_path)

# Read the job description
with open(job_description_path, 'r') as job_description_file:
    job_description = job_description_file.read()

# Extract keywords from the job description
# You can use more sophisticated keyword extraction techniques if needed
keywords = re.findall(r'\b\w+\b', job_description.lower())

# Customize the cover letter by replacing placeholders with keywords
for keyword in keywords:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(f'{{{{Keyword}}}}', keyword)

# Read your resume using python-docx
resume_doc = Document(resume_path)

# Extract relevant content from your resume
resume_text = ""
for paragraph in resume_doc.paragraphs:
    resume_text += paragraph.text + '\n'

# Reference your resume in the cover letter
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        run.text = run.text.replace('{{resume.docx}}', resume_text)

# Save the customized cover letter
document.save('customized_cover_letter.docx')
